# 导入 os 模块
from docx import Document
import os


# 获取当前 py 文件 所在文件夹
cur_dir = 'd:/test/xlsx'

# 列出 cur_dir 下的所有文件名称
cur_file_list = os.listdir(cur_dir)

# 从 docx 模块中 导入 Document 类

# 创建 doc 对象
document = Document()

# 添加 1级 标题
document.add_heading('工作文件夹清单', 0)

# 将 cur_file_list 中每个元素，写入到 docx 文档中的段落
for cur_file in cur_file_list:
    document.add_paragraph(cur_file)

# 保存 docx 文件
document.save(cur_dir + '/'+'清单.docx')
