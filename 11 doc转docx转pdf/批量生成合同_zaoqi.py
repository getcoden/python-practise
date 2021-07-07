from docx import Document
from openpyxl import load_workbook
# 利用os模块建立文件夹，用于存放生成的合同
import os

# 给定合同模板和汇总表所在的文件夹路径，方便复用
path = r"C:\Users\win\Desktop\合同"

# 结合路径判断生成文件夹，规避程序报错而终止的风险
if not os.path.exists(path + '/' + '全部合同'):
    os.mkdir(path + '/' + '全部合同')

workbook = load_workbook(path + '/' + '合同信息表.xlsx')
sheet = workbook.active

for table_row in range(2, sheet.max_row + 1):
    wordfile = Document(path + '/' + '合同模板.docx')
    for table_col in range(1, sheet.max_column + 1):
        old_text = str(sheet.cell(row=1, column=table_col).value)
        new_text = str(sheet.cell(row=table_row, column=table_col).value)
        if ' ' in new_text:
            new_text = new_text.split()[0]

        # 文档Document - 段落Paragraph - 文字块Run
        all_paragraphs = wordfile.paragraphs
        for paragraph in all_paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace(old_text, new_text)

        # 文档Document - 表格Table - 行Row/列Column - 单元格Cell
        all_tables = wordfile.tables
        for table in all_tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(old_text, new_text)

    # 获取公司名用以生成合同的名称
    company = str(sheet.cell(row=table_row, column=1).value)
    wordfile.save(path + '/' + f'全部合同/{company}合同.docx')
