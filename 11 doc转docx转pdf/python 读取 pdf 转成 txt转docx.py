# python 读取 pdf 转成 txt，再转成 word 并设置加粗字体
# 用绝对路径可以读出来，最好是在jupyter notebook里面运行程序

# 参考：https://www.cnblogs.com/wj-1314/p/9429816.html
#import pyocr
from docx import Document
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
import os.path
import importlib
import sys
import time

importlib.reload(sys)
time1 = time.time()
# print("初始时间为：",time1)


text_path = r'D:\教学工具软件\冰点文库下载器 3.2.10 去广告版\冰点文库下载器 3.2.10 去广告版\kpdf\Python语言：程序设计课程教学改革的理想选择-最新教育资料.pdf'
# text_path = r'photo-words.pdf'


def parse():
    '''解析PDF文本，并保存到TXT文件中'''
    fp = open(text_path, 'rb')
    # 用文件对象创建一个PDF文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器，与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码，如果没有密码，就创建一个空的字符串
    doc.initialize()

    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDF，资源管理器，来共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释其对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            # 想要获取文本就获得对象的text属性，
            for x in layout:
                if(isinstance(x, LTTextBoxHorizontal)):
                    with open(r'results.txt', 'a', encoding='utf-8') as f:
                        results = x.get_text()
                        print(results)
                        if '标准答案' in results:
                            continue
                        #import pdb;pdb.set_trace()
                        f.write(results + "\n")


if __name__ == '__main__':
    parse()
    time2 = time.time()
    print("总共消耗时间为:", time2-time1)


# 再转成docx


document = Document()
paragraph = document.add_paragraph('')

str_lines = open(r'results.txt', encoding='utf-8').readlines()
#import pdb;pdb.set_trace()
for line in str_lines:
    if len(line.strip()) < 1:
        continue
    paragraph = document.add_paragraph('')
    if ('A' in line) or ('B' in line) or ('C' in line) or ('D' in line):
        paragraph.add_run(line.strip())
    else:
        paragraph.add_run(line.strip()).bold = True
document.save('results.docx')
