from docx import Document
from docx.shared import RGBColor

obj = Document(r"C:\Users\win\Desktop\xx\告知书.docx")
word = '校园'


def set_run(run):
    # 设置 run 的字体大小、是否加粗以及字体颜色
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color


for p in obj.paragraphs:
    for r in p.runs:
        if word not in r.text:
            pass
        # 获取当前 run 的字体属性
        font_size = r.font.size
        bold = r.bold
        color = r.font.color.rgb
        # 使用关键词切分当前 run 的文本
        rest = r.text.split(word)
        # 清除当前 run 的内容
        r.text = ''
        for text in rest[:-1]:
            run = p.add_run(text=text)
            set_run(run)
            run = p.add_run(word)
            run.font.size = font_size
            run.bold = bold
            run.font.color.rgb = RGBColor(255, 0, 0)
        run = p.add_run(rest[-1])
        set_run(run)
obj.save(r"C:\Users\win\Desktop\xx\处理后.docx")
